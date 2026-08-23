"""
Generate roles-permisos-design.docx from markdown + embedded images + .mmd references.
This ensures the developer gets a readable Word doc while the AI can read .mmd files separately.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BASE = r"D:\CODIGOS\PassPlat\Docs"
DIAGRAMS = os.path.join(BASE, "diagrams")
OUTPUT = os.path.join(BASE, "roles-permisos-design.docx")

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Roles, Permisos y Accesos', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Análisis y Propuesta de UI/UX')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()  # spacer

# ============================================================
# SECCIÓN 1: MODELO DE DATOS
# ============================================================
doc.add_heading('1. Modelo de Datos (PASSWORDS.sql)', level=1)

doc.add_heading('1.1 Tablas Principales', level=2)

# Table: main tables
table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Tabla', 'Descripción', 'Relación']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True

data = [
    ('Roles', 'Catálogo de roles (tenant-scoped o global)', 'PK: Id (int)'),
    ('Permisos', 'Catálogo global de permisos por módulo', 'PK: Id (int)'),
    ('RolesPermisos', 'Pivot many-to-many: Rol ↔ Permiso', 'PK: Id, FK: IdRol, IdPermiso'),
    ('RolesPoliticasPwd', 'Asigna Política de Contraseña a un Rol por Tenant', 'PK: Id, FK: IdRol, IdPolitica, IdTenant'),
    ('Accesos', 'Asigna un Rol a un Usuario para una App + Tenant', 'PK: Id, FK: IdUsuario, IdTenant, IdApp, IdRol'),
    ('Usuarios', 'Usuarios del sistema', 'FK: IdRol (legacy, nullable) → Roles'),
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

doc.add_paragraph()

doc.add_heading('1.2 Relaciones', level=2)
doc.add_paragraph('Usuarios ──── 1:N ────► Accesos ──── N:1 ────► Roles ──── N:M ────► Permisos')
doc.add_paragraph('                        (IdUsuario)            (IdRol)              (via RolesPermisos)')
doc.add_paragraph('                                                   │')
doc.add_paragraph('                                                   └─── N:M ────► PoliticasPwd')
doc.add_paragraph('                                                                   (via RolesPoliticasPwd)')

doc.add_heading('1.3 Índices Únicos', level=2)
table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Light Grid Accent 1'
for i, h in enumerate(['Tabla', 'Índice', 'Columnas', 'Filter']):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
idx_data = [
    ('RolesPermisos', 'UX_RolesPermisos_Activo', '(IdRol, IdPermiso)', 'WHERE Activo = 1'),
    ('RolesPoliticasPwd', 'UX_RolesPol_Tenant_Activo', '(IdTenant, IdRol)', 'WHERE Activo = 1'),
    ('Accesos', 'UX_Accesos_UsuarioApp', '(IdUsuario, IdApp, IdRol)', '—'),
]
for r, row_data in enumerate(idx_data, 1):
    for c, val in enumerate(row_data):
        table2.rows[r].cells[c].text = val

doc.add_heading('1.4 Constraints de Unicidad', level=2)
constraints = [
    'Un usuario puede tener múltiples roles por app (diferente rol por tenant/app)',
    'Un rol solo puede tener UNA política de contraseña activa por tenant',
    'Un par (rol, permiso) solo puede existir una vez activo por rol',
]
for c in constraints:
    doc.add_paragraph(c, style='List Bullet')

# ============================================================
# DIAGRAMA ER (imagen embebida)
# ============================================================
doc.add_heading('Diagrama ER — Modelo de Datos', level=2)
er_img = os.path.join(DIAGRAMS, "er-roles-permisos.png")
if os.path.exists(er_img):
    doc.add_picture(er_img, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('Figura 1: Diagrama Entidad-Relación (er-roles-permisos.mmd)')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
ref = doc.add_paragraph()
run = ref.add_run('📎 Archivo fuente Mermaid: ')
run.font.size = Pt(9)
run.font.italic = True
run = ref.add_run('diagrams/er-roles-permisos.mmd')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x59, 0x4A, 0xE1)
run.font.italic = True

# ============================================================
# SECCIÓN 2: BACKEND EXISTENTE
# ============================================================
doc.add_page_break()
doc.add_heading('2. Backend Existente', level=1)

doc.add_heading('2.1 API Endpoints', level=2)

doc.add_heading('Roles (RolesController)', level=3)
t = doc.add_table(rows=5, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Método', 'Ruta', 'Descripción']):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True
roles_api = [
    ('GET', '/api/roles', 'Listar todos los roles'),
    ('POST', '/api/roles', 'Crear rol (via SP)'),
    ('PUT', '/api/roles/{id}', 'Actualizar rol (via SP)'),
    ('DELETE', '/api/roles/{id}', 'Desactivar rol (via SP)'),
]
for r, row_data in enumerate(roles_api, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

doc.add_paragraph()

doc.add_heading('Permisos (PermisosController)', level=3)
t = doc.add_table(rows=7, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Método', 'Ruta', 'Descripción']):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True
perm_api = [
    ('GET', '/api/permisos/activos', 'Listar permisos activos'),
    ('GET', '/api/permisos/rol/{idRol}', 'Permisos asignados a un rol'),
    ('POST', '/api/permisos', 'Crear permiso'),
    ('POST', '/api/permisos/rol', 'Asignar permiso a rol'),
    ('DELETE', '/api/permisos/rol/{idRol}/{idPermiso}', 'Desasignar permiso de rol'),
    ('DELETE', '/api/permisos/{id}', 'Desactivar permiso'),
]
for r, row_data in enumerate(perm_api, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

doc.add_paragraph()

doc.add_heading('Accesos (AccesosController)', level=3)
t = doc.add_table(rows=6, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Método', 'Ruta', 'Descripción']):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True
accesos_api = [
    ('GET', '/api/accesos/usuario/{idUsuario}', 'Accesos de un usuario'),
    ('GET', '/api/accesos/rol/{idRol}', 'Usuarios con un rol'),
    ('GET', '/api/accesos/tenant-app/{idApp}', 'Accesos por tenant+app'),
    ('POST', '/api/accesos/asignar', 'Asignar rol a usuario'),
    ('POST', '/api/accesos/revocar/{idUsuario}/{idApp}', 'Revocar acceso'),
]
for r, row_data in enumerate(accesos_api, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

doc.add_paragraph()

doc.add_heading('RolesPoliticasPwd (RolesPoliticasPwdController)', level=3)
t = doc.add_table(rows=4, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Método', 'Ruta', 'Descripción']):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True
rpp_api = [
    ('GET', '/api/rolespoliticaspwd/rol/{idRol}', 'Políticas de un rol'),
    ('POST', '/api/rolespoliticaspwd', 'Asignar política a rol'),
    ('POST', '/api/rolespoliticaspwd/{id}/desactivar', 'Desasignar política'),
]
for r, row_data in enumerate(rpp_api, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

doc.add_paragraph()

doc.add_heading('2.2 Servicios', level=2)
t = doc.add_table(rows=5, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Servicio', 'Operaciones', 'Patrón']):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True
svc_data = [
    ('IRolService', 'CRUD + ObtenerTodos/PorTenant/Globales/PorCodigo/Paginado', 'SP-based'),
    ('IPermisoService', 'Crear/Eliminar + ObtenerActivos/PorModulo', 'Repository-based'),
    ('IRolPermisoService', 'Asignar/Desasignar/ObtenerPermisosPorRol', 'Repository-based'),
    ('IAccesoService', 'AsignarAcceso/Revocar/ObtenerAccesosPorRol/Usuario/TenantApp', 'SP-based'),
]
for r, row_data in enumerate(svc_data, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

# ============================================================
# DIAGRAMA ARQUITECTURA (imagen embebida)
# ============================================================
doc.add_heading('Diagrama Arquitectura — Capas del Sistema', level=2)
arch_img = os.path.join(DIAGRAMS, "architecture-roles-permisos.png")
if os.path.exists(arch_img):
    doc.add_picture(arch_img, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('Figura 2: Arquitectura Frontend → API → Servicios → BD (architecture-roles-permisos.mmd)')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

ref = doc.add_paragraph()
run = ref.add_run('📎 Archivo fuente Mermaid: ')
run.font.size = Pt(9)
run.font.italic = True
run = ref.add_run('diagrams/architecture-roles-permisos.mmd')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x59, 0x4A, 0xE1)
run.font.italic = True

# ============================================================
# SECCIÓN 3: UI ACTUAL
# ============================================================
doc.add_page_break()
doc.add_heading('3. UI Actual (Index.razor en /admin/roles-permisos)', level=1)

doc.add_heading('3.1 Estructura Actual', level=2)
structure = [
    'Breadcrumb → Panel Principal → Administración → Roles y Permisos',
    'Stat Cards: Roles | Activos | Permisos | Políticas',
    'Toolbar: Roles y Permisos | Refresh | Nuevo Rol | Nuevo Permiso',
    'SECCIÓN ROLES: tabla con filtros, búsqueda, paginación, acciones (editar, desactivar)',
    'DETALLE DEL ROL SELECCIONADO (inline):',
    '  Tab Permisos: lista agrupada por módulo, toggle switches individuales',
    '  Tab Política: chip + detalles de complejidad + asignar/desasignar',
]
for s in structure:
    doc.add_paragraph(s, style='List Bullet' if not s.startswith('  ') else 'List Bullet 2')

doc.add_heading('3.2 Funcionalidades Existentes', level=2)
t = doc.add_table(rows=15, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Funcionalidad', 'Estado', 'Detalle']):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True
features = [
    ('Listar roles', '✅', 'Con filtros, búsqueda, paginación'),
    ('Crear rol', '✅', 'RolDialog.razor (MaxWidth.Small)'),
    ('Editar rol', '✅', 'RolDialog.razor (MaxWidth.Small)'),
    ('Desactivar rol', '✅', 'Confirm dialog'),
    ('Ver permisos del rol', '✅', 'Agrupados por módulo, toggle switches'),
    ('Asignar permiso a rol', '✅', 'Toggle switch individual'),
    ('Desasignar permiso de rol', '✅', 'Toggle switch individual'),
    ('Crear permiso', '✅', 'PermisoDialog.razor (MaxWidth.Small)'),
    ('Eliminar permiso', '✅', 'Confirm dialog'),
    ('Ver política asignada', '✅', 'Chip + detalles de complejidad'),
    ('Asignar política a rol', '✅', 'RolPoliticaPwdDialog.razor'),
    ('Desasignar política de rol', '✅', 'Confirm dialog'),
    ('Asignar rol a usuario', '❌', 'No existe en esta página'),
    ('Editar permiso', '❌', 'Solo crear, no editar'),
]
for r, row_data in enumerate(features, 1):
    for c, val in enumerate(row_data):
        cell = t.rows[r].cells[c]
        cell.text = val

# ============================================================
# SECCIÓN 4: GAPS
# ============================================================
doc.add_heading('4. Gaps Identificados', level=1)

doc.add_heading('4.1 Gap Crítico: Asignación Usuario ↔ Rol', level=2)
doc.add_paragraph(
    'Problema actual: La asignación de roles a usuarios se maneja en la página de Accesos (/accesos), '
    'NO en la página de Roles/Permisos. Un admin que gestiona roles no puede ver qué usuarios tienen '
    'un rol asignado.'
)
doc.add_paragraph('Flujo actual fragmentado:')
steps = [
    'Admin crea rol en /admin/roles-permisos',
    'Admin asigna permisos al rol en /admin/roles-permisos',
    'Admin debe ir a /accesos para asignar el rol a un usuario',
    'No hay forma de ver desde el rol qué usuarios lo tienen asignado',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('4.2 Gap Medio: Edición de Permisos', level=2)
doc.add_paragraph(
    'El PermisoDialog.razor solo permite CREAR permisos. No hay forma de editar el código, nombre, '
    'descripción o módulo de un permiso existente. Si un admin comete un error al crear un permiso, '
    'debe desactivarlo y crear uno nuevo.'
)

doc.add_heading('4.3 Gap Menor: UX de Permisos', level=2)
minor_gaps = [
    'No hay "Seleccionar todos" / "Deseleccionar todos" por módulo',
    'Los diálogos son estrechos (MaxWidth.Small)',
    'No hay confirmación antes de toggle (acción instantánea)',
    'No hay indicador visual de cuántos usuarios tiene un rol',
]
for g in minor_gaps:
    doc.add_paragraph(g, style='List Bullet')

# ============================================================
# SECCIÓN 5: PROPUESTA UI/UX
# ============================================================
doc.add_page_break()
doc.add_heading('5. Propuesta de UI/UX', level=1)

doc.add_heading('5.1 Layout Propuesto: Tres Pestañas', level=2)
doc.add_paragraph('Roles → Permisos (nueva vista independiente) → Asignación de Usuarios (NUEVO)')

doc.add_heading('5.2 Pestaña 1: Roles (existente, mejorada)', level=2)
improvements = [
    'Wider dialog: CrearRolDialog → MaxWidth.Medium (era Small)',
    'Inline users count: Badge en cada fila mostrando cuántos usuarios tienen el rol',
    'Quick action: Click en rol muestra detalle inline (actual)',
    'Nueva pestaña "Usuarios Asignados" en el detalle del rol',
]
for i in improvements:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('5.3 Pestaña 2: Permisos (nueva vista independiente)', level=2)
doc.add_paragraph(
    'Lista completa de permisos agrupados por módulo, con acciones CRUD. '
    'Incluye: buscar, filtrar por módulo, editar permiso existente, desactivar.'
)

doc.add_heading('5.4 Pestaña 3: Asignación de Usuarios (nueva)', level=2)
doc.add_paragraph(
    'Desde la perspectiva del ROL, ver qué usuarios lo tienen asignado y gestionar asignaciones. '
    'Incluye: dropdown de selección de rol, tabla de usuarios, dialog para asignar usuario a rol.'
)

doc.add_heading('5.5 Detalle del Rol (mejorado)', level=2)
detail_improvements = [
    'Pestaña Permisos: "Seleccionar todos" / "Deselect all" por módulo, contadores por módulo',
    'Nueva pestaña "Usuarios Asignados": lista de usuarios (via GET /api/accesos/rol/{idRol})',
    'Botón para asignar/desasignar usuario desde el detalle del rol',
]
for i in detail_improvements:
    doc.add_paragraph(i, style='List Bullet')

# ============================================================
# DIAGRAMA DE FLUJO (imagen embebida)
# ============================================================
doc.add_heading('Diagrama de Flujo — Workflow Completo', level=2)
flow_img = os.path.join(DIAGRAMS, "flow-roles-permisos.png")
if os.path.exists(flow_img):
    doc.add_picture(flow_img, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('Figura 3: Flujo de usuario con 5 subprocesos (flow-roles-permisos.mmd)')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

ref = doc.add_paragraph()
run = ref.add_run('📎 Archivo fuente Mermaid: ')
run.font.size = Pt(9)
run.font.italic = True
run = ref.add_run('diagrams/flow-roles-permisos.mmd')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x59, 0x4A, 0xE1)
run.font.italic = True

# ============================================================
# SECCIÓN 6: CAMBIOS BACKEND
# ============================================================
doc.add_page_break()
doc.add_heading('6. Cambios Requeridos en Backend', level=1)

doc.add_heading('6.1 Endpoints Nuevos Necesarios', level=2)
t = doc.add_table(rows=4, cols=4)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Endpoint', 'Método', 'Descripción', 'Prioridad']):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True
new_eps = [
    ('/api/permisos/{id}', 'GET', 'Obtener permiso por ID', 'Alta'),
    ('/api/permisos/{id}', 'PUT', 'Actualizar permiso', 'Alta'),
    ('/api/roles/{id}/usuarios', 'GET', 'Usuarios con un rol', 'Media'),
]
for r, row_data in enumerate(new_eps, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

doc.add_heading('6.2 Endpoints Existentes Reutilizables', level=2)
reusable = [
    'GET /api/accesos/rol/{idRol} → Usuarios con un rol ✅',
    'POST /api/accesos/asignar → Asignar rol a usuario ✅',
    'POST /api/accesos/revocar/{idUsuario}/{idApp} → Revocar acceso ✅',
    'GET /api/permisos/activos → Todos los permisos activos ✅',
    'GET /api/permisos/rol/{idRol} → Permisos de un rol ✅',
]
for r in reusable:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('6.3 Servicios a Implementar', level=2)
doc.add_paragraph('PermisoService — agregar métodos faltantes:')
methods = [
    'Task<Result<PermisoDto?>> ObtenerPorIdAsync(int id, CancellationToken ct = default);',
    'Task<Result<PermisoDto>> ActualizarAsync(int id, CrearPermisoDto dto, CancellationToken ct = default);',
]
for m in methods:
    p = doc.add_paragraph()
    run = p.add_run(m)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

# ============================================================
# SECCIÓN 7: FLUJO USUARIO
# ============================================================
doc.add_heading('7. Flujo de Usuario Propuesto', level=1)

doc.add_heading('7.1 Asignar permisos a un rol', level=2)
steps_71 = [
    'Admin → /admin/roles-permisos → pestaña "Roles"',
    'Selecciona rol "Administrador"',
    'En detalle, pestaña "Permisos" → ve todos los permisos agrupados',
    'Marca/desmarca permisos (toggles existentes)',
    'Opcional: "Seleccionar todos" del módulo "Seguridad"',
]
for i, s in enumerate(steps_71, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('7.2 Asignar rol a un usuario', level=2)
doc.add_paragraph('Opción A (desde RolesPermisos):', style='List Bullet')
steps_a = [
    'Admin → /admin/roles-permisos → pestaña "Asignación de Usuarios"',
    'Selecciona rol "Administrador"',
    'Ve lista de usuarios asignados',
    'Click "+ Asignar Usuario" → dialog con búsqueda de usuario + selección de App',
    'Confirma → usuario aparece en la lista',
]
for i, s in enumerate(steps_a, 1):
    doc.add_paragraph(f'  {i}. {s}')

doc.add_paragraph('Opción B (desde Usuarios):', style='List Bullet')
steps_b = [
    'Admin → /usuarios → selecciona usuario → pestaña "Accesos"',
    'Click "+ Asignar Acceso" → dialog existente',
    'Selecciona App + Rol → confirma',
]
for i, s in enumerate(steps_b, 1):
    doc.add_paragraph(f'  {i}. {s}')

doc.add_heading('7.3 Crear y configurar un nuevo rol completo', level=2)
steps_73 = [
    'Admin → /admin/roles-permisos → "Nuevo Rol"',
    'Completa formulario → crea rol',
    'Selecciona el nuevo rol en la tabla',
    'Pestaña "Permisos" → asigna permisos',
    'Pestaña "Política" → asigna política de contraseña',
    'Pestaña "Asignación de Usuarios" (nueva) → asigna usuarios',
]
for i, s in enumerate(steps_73, 1):
    doc.add_paragraph(f'{i}. {s}')

# ============================================================
# SECCIÓN 8: ESTIMACIÓN
# ============================================================
doc.add_heading('8. Estimación de Esfuerzo', level=1)
t = doc.add_table(rows=8, cols=3)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Componente', 'Esfuerzo', 'Detalle']):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True
effort = [
    ('PermisoDialog modo edición', 'Bajo', 'Agregar IsEdit + Permiso param, PUT endpoint'),
    ('"Select all" por módulo', 'Bajo', 'Lógica en PermisoGrupo'),
    ('Pestaña "Asignación de Usuarios"', 'Medio', 'Nuevo tab + list + dialog'),
    ('Dialog "Asignar Usuario a Rol"', 'Medio', 'Búsqueda de usuario + selección App'),
    ('Wider dialogs (MaxWidth.Medium)', 'Trivial', 'Cambiar 1 línea en cada dialog'),
    ('Backend: GET /api/permisos/{id}', 'Bajo', 'Nuevo endpoint'),
    ('Backend: PUT /api/permisos/{id}', 'Bajo', 'Nuevo endpoint + servicio'),
]
for r, row_data in enumerate(effort, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

doc.add_paragraph()
total = doc.add_paragraph()
run = total.add_run('Total estimado: ~4-6 horas')
run.font.bold = True
run.font.size = Pt(12)

# ============================================================
# SECCIÓN 9: RECOMENDACIONES
# ============================================================
doc.add_heading('9. Recomendaciones', level=1)
recs = [
    'Priorizar la pestaña "Asignación de Usuarios" — Es el gap más crítico.',
    'Agregar modo edición a PermisoDialog — Cambio bajo que elimina desactivar/recrear.',
    'Wider dialogs — Cambio trivial que mejora significativamente la experiencia.',
    'No duplicar funcionalidad — Considerar link rápido "Ver usuarios" en detalle del rol vs. mover toda la asignación.',
    'Futuro: drag & drop para reordenar permisos. Baja prioridad.',
]
for i, r in enumerate(recs, 1):
    doc.add_paragraph(f'{i}. {r}')

# ============================================================
# REFERENCIA A ARCHIVOS FUENTE
# ============================================================
doc.add_page_break()
doc.add_heading('Anexo: Archivos Fuente', level=1)
doc.add_paragraph(
    'Este documento fue generado automáticamente. Los diagramas Mermaid fuente (.mmd) '
    'se encuentran en la carpeta diagrams/ y pueden re-renderizarse con:'
)
p = doc.add_paragraph()
run = p.add_run('npx @mermaid-js/mermaid-cli -i <input.mmd> -o <output.png> -w 1600 -H 1000 -b white')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Archivos incluidos:')
files_list = [
    'diagrams/er-roles-permisos.mmd → Figura 1 (Diagrama ER)',
    'diagrams/architecture-roles-permisos.mmd → Figura 2 (Arquitectura)',
    'diagrams/flow-roles-permisos.mmd → Figura 3 (Flujo de usuario)',
    'diagrams/prototype-roles-permisos.html → Prototipo interactivo (abrir en navegador)',
]
for f in files_list:
    doc.add_paragraph(f, style='List Bullet')

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f"Generado: {OUTPUT}")
print(f"Tamaño: {os.path.getsize(OUTPUT) / 1024:.0f} KB")
